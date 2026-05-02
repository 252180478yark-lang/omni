"""从上传的文件中提取文本内容。
Excel/CSV: 规则解析（快速）
docx/pdf/txt: 先提取原文，再用大模型总结关键信息
"""

from __future__ import annotations

import csv
import io
import json
import logging
from pathlib import Path
from typing import Any

import httpx

logger = logging.getLogger(__name__)

_AI_HUB_URL = ""
_SUMMARY_MODEL = ""


def _get_ai_config() -> tuple[str, str]:
    global _AI_HUB_URL, _SUMMARY_MODEL
    if not _AI_HUB_URL:
        from app.config import settings
        _AI_HUB_URL = settings.ai_hub_url.rstrip("/")
        _SUMMARY_MODEL = settings.review_model
    return _AI_HUB_URL, _SUMMARY_MODEL


def extract_text(data: bytes, filename: str, context: str = "人群画像") -> str:
    ext = Path(filename or "").suffix.lower()

    if ext == ".csv":
        raw = _parse_csv_text(data)
        if len(raw) > 500:
            return _llm_summarize(raw, context)
        return raw
    if ext in (".xlsx", ".xlsm", ".xls"):
        raw = _parse_xlsx_text(data)
        if raw.startswith("（"):
            return raw
        if len(raw) > 500:
            return _llm_summarize(raw, context)
        return raw
    if ext == ".doc":
        return "（.doc 格式暂不支持，请另存为 .docx 后重新上传）"

    raw_text = ""
    if ext == ".txt":
        raw_text = _decode_text(data)
    elif ext == ".docx":
        raw_text = _parse_docx_text(data)
    elif ext == ".pdf":
        raw_text = _parse_pdf_text(data)

    if not raw_text or raw_text.startswith("（"):
        return raw_text

    return _llm_summarize(raw_text, context)


_SECTION_AUDIENCE_PROFILE = """## 1. 核心人群特征
- 年龄段分布（如 25-35 岁占比最高）
- 性别比例
- 地域分布（一二三线城市占比、TOP 省市）
- 消费能力层级

## 2. 兴趣与行为标签
- 主要兴趣分类（如美妆、护肤、母婴等）
- 消费行为特征（高频复购、价格敏感、品牌忠诚等）
- 内容偏好（短视频类型偏好、互动行为特征）
- 活跃时段

## 3. 人群规模与覆盖
- 人群包总覆盖人数
- 各细分人群的规模
- 与竞品人群的重合度（如有）

## 4. 对素材创作的启示
- 这类人群最可能被什么类型的内容吸引？
- 痛点和需求是什么？
- 价格敏感度如何？怎么做价格锚点？
- 什么样的钩子和卖点最有效？

## 5. 对投放策略的启示
- 建议的定向方式
- 建议的投放时段
- 预算分配建议"""


_SECTION_TARGETING = """## 1. 定向策略概述
- 使用了哪些定向方式（行为兴趣定向 / DMP 人群包 / 达人相似 / 智能定向等）
- 核心定向逻辑是什么（如"竞品粉丝 + 品类兴趣 + 排除已购"）

## 2. 人群包详情
- 每个人群包的构成规则（包含什么标签、排除什么标签）
- 人群包覆盖规模
- 各人群包之间的区别和分工

## 3. 定向标签明细
- 行为标签列表（如"购买过竞品""浏览过护肤内容"等）
- 兴趣标签列表
- 排除标签列表
- 自定义人群条件

## 4. 历史效果数据（如有）
- 各人群包的历史 CTR、CVR、CPM 表现
- 哪个人群包效果最好 / 最差
- 历史优化调整记录

## 5. 对素材匹配与投放的启示
- 这个人群包最匹配什么类型的素材？
- 钩子和卖点应该怎么设计？
- 不同人群包之间的素材分配建议、建议的定向优化方向"""


_SECTION_GENERIC = """## 1. 核心信息（按类别整理）
- 人群特征（如提及）
- 定向策略（如提及）
- 数据亮点（如提及）
- 优化经验（如提及）

## 2. 对素材创作的启示
- 钩子、卖点、场景方向的可参考结论

## 3. 对投放策略的启示
- 定向、人群包、匹配度调整的可参考结论"""


_FILE_SUMMARY_DISCIPLINE = """## 纪律（必须遵守）
1. 文件中没有明确提供的信息，写"文件未提供"。禁止编造数字、标签、指标、人群规模。
2. 表格数据转为易读的文字描述，并指出关键发现（TOP/差异/趋势）。
3. 直接输出整理后的内容，不加前缀说明。"""


def _llm_summarize(raw_text: str, context: str) -> str:
    ai_url, model = _get_ai_config()
    if not ai_url:
        return raw_text

    truncated = raw_text[:12000]

    if "画像" in context:
        role_hint = "人群画像分析"
        section_block = _SECTION_AUDIENCE_PROFILE
    elif "圈包" in context or "定向" in context:
        role_hint = "DMP 人群包定向策略分析"
        section_block = _SECTION_TARGETING
    else:
        role_hint = "投放决策支持信息提取"
        section_block = _SECTION_GENERIC

    prompt = f"""你是一位资深的巨量千川投放优化师，当前任务：{role_hint}。

以下是一份与「{context}」相关的文件。请提取所有对投放决策有价值的信息，按下方结构整理输出。

{_FILE_SUMMARY_DISCIPLINE}

## 输出结构
{section_block}

## 文件原文
{truncated}"""

    # 飞轮：同步拉取累积规则（file_parser 是同步函数,不走 asyncio.run 避免事件循环冲突）
    try:
        from app.config import settings as _settings
        ke_base = _settings.knowledge_engine_url.rstrip("/")
        rule_resp = httpx.post(
            f"{ke_base}/api/v1/prompt/render-rules",
            json={"node_id": "review.file_parser", "scope": {"context": context}},
            timeout=5.0,
        )
        if rule_resp.status_code == 200:
            suffix = ((rule_resp.json() or {}).get("data") or {}).get("suffix") or ""
            if suffix:
                prompt += suffix
    except Exception:
        pass

    try:
        resp = httpx.post(
            f"{ai_url}/v1/chat/completions",
            json={
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1,
                "max_tokens": 2000,
            },
            timeout=60.0,
        )
        if resp.status_code == 200:
            body = resp.json()
            content = body["choices"][0]["message"]["content"]
            return content.strip()
    except Exception as e:
        logger.warning("LLM summarize failed: %s", e)

    return raw_text


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return data.decode("utf-8", errors="replace")


def _parse_csv_text(data: bytes) -> str:
    text = _decode_text(data)
    reader = csv.reader(io.StringIO(text))
    lines = []
    for row in reader:
        lines.append(" | ".join(cell.strip() for cell in row if cell.strip()))
    return "\n".join(lines)


def _parse_xlsx_text(data: bytes) -> str:
    raw = data
    # Handle encrypted Office files (WPS/Excel default encryption)
    if data[:4] == b'\xd0\xcf\x11\xe0':
        try:
            import msoffcrypto
            f_in = io.BytesIO(data)
            f_out = io.BytesIO()
            ms_file = msoffcrypto.OfficeFile(f_in)
            if ms_file.is_encrypted():
                try:
                    ms_file.load_key(password="")
                    ms_file.decrypt(f_out)
                    raw = f_out.getvalue()
                except Exception:
                    return (
                        "（该Excel文件已加密/受保护，无法自动读取。\n"
                        "解决方法：用WPS或Excel打开文件 → 文件 → 另存为 → 选择.xlsx格式 → 取消勾选「加密」→ 保存 → 重新上传）"
                    )
        except ImportError:
            pass

    # Try openpyxl first (xlsx)
    from openpyxl import load_workbook
    try:
        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            if ws.title:
                lines.append(f"[{ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines) if lines else "（Excel 无内容）"
    except Exception:
        pass

    # Fallback to xlrd (xls)
    try:
        import xlrd
        book = xlrd.open_workbook(file_contents=raw)
        lines = []
        for sheet in book.sheets():
            if sheet.name:
                lines.append(f"[{sheet.name}]")
            for rx in range(sheet.nrows):
                cells = [str(sheet.cell_value(rx, cx)).strip() for cx in range(sheet.ncols)
                         if str(sheet.cell_value(rx, cx)).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines) if lines else "（Excel 无内容）"
    except Exception:
        pass

    return "（Excel 文件无法读取，可能文件已加密或格式不兼容，请尝试用WPS/Excel另存为新的.xlsx文件）"


def _parse_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "（缺少 python-docx 库）"

    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        return "（.docx 文件无法读取）"

    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _parse_pdf_text(data: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return "（缺少 PyPDF2 库）"

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception:
        return "（PDF 文件无法读取）"

    lines = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            lines.append(text.strip())
    return "\n".join(lines)
